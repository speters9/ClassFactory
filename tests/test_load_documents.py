
import shutil
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from pyprojroot.here import here

from class_factory.utils.load_documents import LessonLoader

wd = here()


@pytest.fixture
def mock_paths():
    # Create temporary directories
    temp_dirs = {
        "syllabus_path": Path(tempfile.mkdtemp()) / "syllabus.txt",
        "reading_dir": Path(tempfile.mkdtemp()),
        "slide_dir": Path(tempfile.mkdtemp()),
        "output_dir": Path(tempfile.mkdtemp()),
        "project_dir": Path(tempfile.mkdtemp()),
    }
    # Create a dummy syllabus file
    temp_dirs["syllabus_path"].write_text("Syllabus content here.")

    # Yield the paths to the test
    yield temp_dirs

    # Cleanup after the test
    for temp_dir in temp_dirs.values():
        if temp_dir.is_dir():
            shutil.rmtree(temp_dir, ignore_errors=True)


@pytest.fixture
def mock_lesson_loader(mock_paths):
    # Create mock readings within the reading directory
    lesson_dir = mock_paths["reading_dir"] / "L1"
    lesson_dir.mkdir(parents=True)
    (lesson_dir / "reading1.txt").write_text("Reading 1 content")
    (lesson_dir / "reading2.txt").write_text("Reading 2 content")

    return LessonLoader(
        syllabus_path=mock_paths["syllabus_path"],
        reading_dir=mock_paths["reading_dir"],
        slide_dir=mock_paths["slide_dir"],
        project_dir=mock_paths["project_dir"]
    )


def test_initialization_valid_paths(mock_paths):
    syllabus = mock_paths["syllabus_path"]
    reading_dir = mock_paths["reading_dir"]
    loader = LessonLoader(syllabus_path=syllabus, reading_dir=reading_dir)

    assert loader.syllabus_path == syllabus
    assert loader.reading_dir == reading_dir


def test_validate_file_path(tmp_path):
    file_path = tmp_path / "file.txt"
    file_path.touch()  # Create an actual file

    # Positive check
    validated_path = LessonLoader._validate_file_path(file_path, "syllabus file")
    assert validated_path == file_path

    # Negative check
    with pytest.raises(FileNotFoundError):
        LessonLoader._validate_file_path(Path("nonexistent_file.txt"), "syllabus file")


def test_validate_dir_path_create(tmp_path):
    dir_path = tmp_path / "new_dir"
    validated_path = LessonLoader._validate_dir_path(dir_path, "test dir", create_if_missing=True)
    assert validated_path.exists() and validated_path.is_dir()


def test_validate_dir_path_no_create(tmp_path):
    dir_path = tmp_path / "missing_dir"
    with pytest.raises(NotADirectoryError):
        LessonLoader._validate_dir_path(dir_path, "test dir")


def test_load_directory(mock_paths):
    reading_dir = mock_paths["reading_dir"]
    (reading_dir / "L1_sample.pdf").touch()
    (reading_dir / "L1_sample.txt").touch()

    # Mock `load_readings` to simulate content
    def mock_load_readings(file):
        return f"Mocked document content for {file.name}"

    with patch.object(LessonLoader, 'load_readings', side_effect=mock_load_readings):
        loader = LessonLoader(syllabus_path=mock_paths["syllabus_path"], reading_dir=reading_dir)
        documents = loader.load_directory(reading_dir)

        assert len(documents) == 2
        assert "Mocked document content for L1_sample.pdf" in documents
        assert "Mocked document content for L1_sample.txt" in documents


def test_load_lessons(mock_paths):
    reading_dir = mock_paths["reading_dir"]
    lesson_dir = reading_dir / "L1"
    lesson_dir.mkdir()
    (lesson_dir / "L1_sample.txt").touch()

    with patch.object(LessonLoader, '_validate_dir_path', side_effect=lambda path, name, must_contain_contents=False: Path(path)):
        loader = LessonLoader(syllabus_path=mock_paths["syllabus_path"], reading_dir=reading_dir)
        documents = loader.load_lessons(lesson_number_or_range=range(1, 2))
        assert len(documents) == 1


def test_extract_text_from_pdf(mock_paths):
    pdf_path = mock_paths["output_dir"] / "sample.pdf"
    pdf_path.touch()

    mock_reader = MagicMock()
    mock_reader.pages = [MagicMock()]
    mock_reader.pages[0].extract_text.return_value = "Sample text content from mock PDF."

    with patch('pypdf.PdfReader', return_value=mock_reader):
        loader = LessonLoader(syllabus_path=mock_paths["syllabus_path"], reading_dir=mock_paths["reading_dir"])
        text = loader.extract_text_from_pdf(pdf_path)
        assert text == "Sample text content from mock PDF."


def test_convert_pdf_to_docx(mock_paths):
    pdf_path = mock_paths["output_dir"] / "sample.pdf"
    pdf_path.touch()

    with patch('class_factory.utils.load_documents.Converter') as mock_converter:
        mock_instance = mock_converter.return_value
        mock_instance.convert.return_value = None
        mock_instance.close.return_value = None

        loader = LessonLoader(syllabus_path=mock_paths["syllabus_path"], reading_dir=mock_paths["reading_dir"])
        docx_path = loader.convert_pdf_to_docx(pdf_path)

        assert docx_path.suffix == ".docx"
        mock_instance.convert.assert_called_once()
        mock_instance.close.assert_called_once()


def test_extract_lesson_objectives(mock_paths):
    syllabus_path = mock_paths["syllabus_path"]
    doc = Document()
    doc.add_paragraph("Lesson 1:")
    doc.add_paragraph("Objective 1")
    doc.add_paragraph("Objective 2")
    doc.add_paragraph("Lesson 2:")
    doc.add_paragraph("Objective A")
    doc.save(syllabus_path)

    loader = LessonLoader(syllabus_path=syllabus_path, reading_dir=mock_paths["reading_dir"])
    objectives = loader.extract_lesson_objectives(current_lesson=1)

    assert "Objective 1" in objectives
    assert "Objective 2" in objectives
    assert "Objective A" in objectives


def test_find_docx_indices(mock_paths):
    syllabus = ["Introduction", "Lesson 1:", "Objective 1", "Objective 2", "Lesson 2:", "Objective A", "Objective B"]
    with patch('class_factory.utils.load_documents.LessonLoader._validate_dir_path', side_effect=lambda path, name, must_contain_contents=False: Path(path)), \
            patch('pathlib.Path.is_file', return_value=True):

        loader = LessonLoader(syllabus_path="dummy_syllabus.docx", reading_dir=mock_paths['reading_dir'])
        prev, curr, nxt, end = loader.find_docx_indices(syllabus, current_lesson=1)
        assert syllabus[curr] == "Lesson 1:"
        assert syllabus[nxt] == "Lesson 2:"


def test_extract_text_from_pdf_ocr_unavailable(mock_paths):
    pdf_path = mock_paths["output_dir"] / "sample.pdf"
    pdf_path.touch()

    with patch('class_factory.utils.load_documents.LessonLoader.ocr_available', return_value=False), \
            patch('pypdf.PdfReader', side_effect=ImportError("OCR support requires additional packages")):

        loader = LessonLoader(syllabus_path=mock_paths["syllabus_path"], reading_dir=mock_paths["reading_dir"])

        with pytest.raises(ImportError, match="OCR support requires additional packages"):
            loader.extract_text_from_pdf(pdf_path)


def test_load_beamer_presentation(mock_paths):
    tex_path = mock_paths["slide_dir"] / "sample.tex"
    tex_path.write_text("Beamer slide content")

    loader = LessonLoader(syllabus_path=mock_paths["syllabus_path"], reading_dir=mock_paths["reading_dir"])
    content = loader.load_beamer_presentation(tex_path)

    assert content == "Beamer slide content"


def test_load_readings_logging_warning(mock_paths, caplog):
    unreadable_file = mock_paths["reading_dir"] / "unreadable.pdf"
    unreadable_file.touch()

    loader = LessonLoader(syllabus_path=mock_paths["syllabus_path"], reading_dir=mock_paths["reading_dir"])

    with patch.object(loader, 'extract_text_from_pdf', return_value=""):
        loader.load_readings(unreadable_file)

    assert "No readable text found" in caplog.text


@pytest.fixture
def mock_reading_structure():
    # Create a temporary directory structure for reading_dir
    with tempfile.TemporaryDirectory() as tmpdir:
        base_path = Path(tmpdir)
        # Create lesson directories that match and don't match the pattern
        (base_path / "L1").mkdir()
        (base_path / "Lesson_2").mkdir()
        (base_path / "Lecture3").mkdir()
        (base_path / "Week4").mkdir()
        (base_path / "non_lesson_dir").mkdir()  # Should not be processed
        (base_path / "Lesson_5").mkdir()  # Out of specified range for test
        # Populate each directory with a sample file
        for lesson_dir in base_path.iterdir():
            if lesson_dir.is_dir():
                (lesson_dir / f"{lesson_dir.name}_sample.txt").write_text("Sample reading content.")
        yield base_path


def test_load_lessons(mock_reading_structure, mock_paths):
    # Mock paths and initialize LessonLoader with the mock reading structure
    reading_dir = mock_reading_structure
    syllabus_path = mock_paths['syllabus_path']  # Not used in this test

    with patch('class_factory.utils.load_documents.LessonLoader._validate_dir_path', side_effect=lambda path, name, must_contain_contents=False: Path(path)):
        with pytest.raises(ValueError, match="Directory structure validation failed with the following issues:"):
            loader = LessonLoader(syllabus_path=syllabus_path, reading_dir=reading_dir)

            # Define the range of lessons to load (testing range 1-4)
            lesson_range = range(1, 5)
            loaded_lessons = loader.load_lessons(lesson_range)

            # Check that only directories within the range are loaded
            assert "1" in loaded_lessons  # Matches L1
            assert "2" in loaded_lessons  # Matches Lesson_2
            assert "3" in loaded_lessons  # Matches Lecture3
            assert "4" in loaded_lessons  # Matches Week4
            assert "5" not in loaded_lessons  # Out of specified range
            assert "non_lesson_dir" not in loaded_lessons  # Does not match pattern

            # Verify the contents of each loaded lesson
            for lesson, readings in loaded_lessons.items():
                assert len(readings) == 1  # One file per lesson directory
                assert readings[0] == f"{lesson}_sample.txt"  # Check if the sample file is loaded correctly


if __name__ == "__main__":
    pytest.main([__file__])
